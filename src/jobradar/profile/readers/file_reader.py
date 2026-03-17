"""File and URL CV reader — supports .md/.txt/.pdf/.docx and http(s):// URLs."""

from __future__ import annotations

import logging
import re
from pathlib import Path
from urllib.parse import urlparse

from .base import ProfileReader

logger = logging.getLogger(__name__)


class FileAndUrlReader(ProfileReader):
    """Handles local files (.md, .txt, .pdf, .docx) and remote URLs."""

    def can_handle(self, source: str) -> bool:
        return True  # fallback reader — handles everything

    def read(self, source: str) -> str:
        s = source.strip()
        if s.startswith("http://") or s.startswith("https://"):
            return self._read_url(s)
        return self._read_file(Path(s))

    # ── Local files ────────────────────────────────────────────────

    def _read_file(self, path: Path) -> str:
        if not path.exists():
            raise FileNotFoundError(f"CV not found: {path}")
        suffix = path.suffix.lower()
        if suffix in (".md", ".txt", ""):
            return path.read_text(encoding="utf-8")
        elif suffix == ".pdf":
            return self._parse_pdf(path.read_bytes(), str(path))
        elif suffix in (".docx", ".doc"):
            return self._parse_docx(path.read_bytes(), str(path))
        else:
            raise ValueError(f"Unsupported format: {suffix}. Use .md .txt .pdf .docx")

    # ── Remote URLs ────────────────────────────────────────────────

    def _read_url(self, url: str) -> str:
        import httpx
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 Chrome/124.0 Safari/537.36"
            ),
            "Accept": "text/html,application/pdf,*/*",
        }
        try:
            with httpx.Client(follow_redirects=True, timeout=30.0) as client:
                resp = client.get(url, headers=headers)
                resp.raise_for_status()
        except Exception as e:
            if "SSL" in str(e).upper() or "CERTIFICATE" in str(e).upper():
                with httpx.Client(follow_redirects=True, timeout=30.0, verify=False) as client:
                    resp = client.get(url, headers=headers)
                    resp.raise_for_status()
            else:
                raise

        ct = resp.headers.get("content-type", "").lower().split(";")[0].strip()
        ext = Path(urlparse(url).path).suffix.lower()
        fmt = self._detect_format(ct, ext, url)
        logger.info("CV URL format: %s  (content-type=%s)", fmt, ct)

        if fmt == "pdf":
            return self._parse_pdf(resp.content, url)
        elif fmt == "docx":
            return self._parse_docx(resp.content, url)
        elif fmt == "html":
            return self._extract_html(resp.text, url)
        else:
            return resp.text  # markdown or plain text

    def _detect_format(self, ct: str, ext: str, url: str) -> str:
        if "text/html" in ct:
            return "html"
        if "application/pdf" in ct:
            return "pdf"
        if "wordprocessingml" in ct:
            return "docx"
        ext_map = {".pdf": "pdf", ".docx": "docx", ".doc": "docx",
                   ".md": "markdown", ".txt": "text", ".html": "html", ".htm": "html"}
        if ext in ext_map:
            return ext_map[ext]
        if "raw.githubusercontent.com" in url or "gist.github" in url:
            return "markdown"
        return "text"

    # ── Format parsers ─────────────────────────────────────────────

    def _parse_pdf(self, data: bytes, source: str) -> str:
        try:
            from pypdf import PdfReader
            import io
        except ImportError:
            raise ImportError("pypdf required: pip install pypdf")
        reader = PdfReader(io.BytesIO(data))
        pages = [p.extract_text() for p in reader.pages if p.extract_text()]
        text = "\n\n".join(pages)
        logger.info("PDF '%s': %d chars from %d pages", source, len(text), len(pages))
        return text

    def _parse_docx(self, data: bytes, source: str) -> str:
        try:
            import docx, io
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")
        doc = docx.Document(io.BytesIO(data))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        logger.info("DOCX '%s': %d chars", source, len(text))
        return text

    def _extract_html(self, html: str, url: str) -> str:
        try:
            import html2text
            h = html2text.HTML2Text()
            h.ignore_links = True
            h.ignore_images = True
            h.body_width = 0
            return _clean(h.handle(html))
        except ImportError:
            pass
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            return _clean(soup.get_text(separator="\n"))
        except ImportError:
            pass
        return _clean(re.sub(r"<[^>]+>", " ", html))


def _clean(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    return re.sub(r" {2,}", " ", text).strip()
