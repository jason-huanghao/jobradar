"""CV reader — supports local files AND remote URLs.

Supported sources:
  Local files   .md / .txt / .pdf / .docx / .doc
  Remote URLs   http(s):// → auto-detects format from Content-Type + extension
                Handles HTML, PDF, Markdown, plain text, DOCX

Usage:
    from .cv_reader import read_cv_file, read_cv_from_url, read_cv_auto

    # Local file
    text = read_cv_file(Path("cv/cv.md"))

    # Remote URL (any supported format)
    text = read_cv_from_url("https://example.com/cv.pdf")

    # Auto (string: URL or path)
    text = read_cv_auto("https://example.com/resume.html")
    text = read_cv_auto("./cv/cv.md")
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from urllib.parse import urlparse

logger = logging.getLogger(__name__)

# ── Public API ─────────────────────────────────────────────────────


def read_cv_auto(source: str) -> str:
    """Read a CV from a local path or a URL — auto-detected.

    Args:
        source: A file path (str or Path) or an http(s):// URL.

    Returns:
        Plain text suitable for LLM processing.
    """
    s = str(source).strip()
    if s.startswith("http://") or s.startswith("https://"):
        return read_cv_from_url(s)
    return read_cv_file(Path(s))


def read_cv_file(cv_path: Path) -> str:
    """Read a CV from a local file and return its text content.

    Supports: .md .txt .pdf .docx .doc
    """
    cv_path = Path(cv_path)
    if not cv_path.exists():
        raise FileNotFoundError(f"CV not found: {cv_path}")

    suffix = cv_path.suffix.lower()
    if suffix in (".md", ".txt", ""):
        return _read_text(cv_path)
    elif suffix == ".pdf":
        return _read_pdf_bytes(cv_path.read_bytes(), cv_path.name)
    elif suffix in (".docx", ".doc"):
        return _read_docx_bytes(cv_path.read_bytes(), cv_path.name)
    else:
        raise ValueError(
            f"Unsupported CV format: '{suffix}'. "
            "Supported: .md .txt .pdf .docx .doc  — or pass a URL."
        )


def read_cv_from_url(url: str) -> str:
    """Download a CV from a URL and return its text content.

    Supports:
        - HTML pages  (extracts main content, strips nav/ads)
        - PDF         (downloads then extracts text)
        - Markdown    (returns raw text)
        - DOCX        (downloads then extracts text)
        - Plain text  (returns as-is)

    The format is detected from:
        1. HTTP Content-Type header (most reliable)
        2. URL path extension (fallback)

    Raises:
        ImportError  — if a required library is missing
        ValueError   — if the URL format is unsupported
        httpx.*      — on network errors
    """
    import httpx

    logger.info("Fetching CV from URL: %s", url)
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
        ),
        "Accept": "text/html,application/pdf,application/octet-stream,*/*",
    }

    def _fetch(verify_ssl: bool) -> httpx.Response:
        with httpx.Client(follow_redirects=True, timeout=30.0, verify=verify_ssl) as client:
            resp = client.get(url, headers=headers)
            resp.raise_for_status()
            return resp

    try:
        resp = _fetch(verify_ssl=True)
    except Exception as ssl_err:
        if "CERTIFICATE" in str(ssl_err).upper() or "SSL" in str(ssl_err).upper():
            logger.warning("SSL verification failed for %s — retrying without verification", url)
            resp = _fetch(verify_ssl=False)
        else:
            raise

    content_type = resp.headers.get("content-type", "").lower().split(";")[0].strip()
    url_path = urlparse(url).path.lower()
    ext = Path(url_path).suffix

    # Detect format
    fmt = _detect_url_format(content_type, ext, url)
    logger.info("URL CV format detected: %s (content-type=%s, ext=%s)", fmt, content_type, ext)

    if fmt == "html":
        return _extract_html_text(resp.text, url)
    elif fmt == "pdf":
        return _read_pdf_bytes(resp.content, url)
    elif fmt in ("markdown", "text"):
        return resp.text
    elif fmt == "docx":
        return _read_docx_bytes(resp.content, url)
    else:
        # Best-effort: try as plain text
        logger.warning("Unknown format '%s' for URL %s — treating as plain text", fmt, url)
        return resp.text


# ── Format detection ───────────────────────────────────────────────


def _detect_url_format(content_type: str, ext: str, url: str) -> str:
    """Return one of: html | pdf | markdown | text | docx | unknown"""
    # Content-Type takes priority
    if "text/html" in content_type:
        return "html"
    if "application/pdf" in content_type:
        return "pdf"
    if "application/vnd.openxmlformats-officedocument.wordprocessingml" in content_type:
        return "docx"
    if content_type in ("text/markdown", "text/x-markdown"):
        return "markdown"
    if content_type in ("text/plain",):
        return "text" if ext not in (".md",) else "markdown"

    # Fallback: extension
    ext_map = {
        ".pdf":  "pdf",
        ".docx": "docx",
        ".doc":  "docx",
        ".md":   "markdown",
        ".txt":  "text",
        ".html": "html",
        ".htm":  "html",
    }
    if ext in ext_map:
        return ext_map[ext]

    # GitHub raw URLs: if path ends with known extension it's often text/plain
    if "raw.githubusercontent.com" in url or "gist.githubusercontent.com" in url:
        return "markdown" if ".md" in url else "text"

    return "unknown"


# ── Format-specific readers ────────────────────────────────────────


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _read_pdf_bytes(data: bytes, source: str) -> str:
    try:
        from pypdf import PdfReader
        import io
    except ImportError:
        raise ImportError("pypdf is required: pip install pypdf")

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())

    if not pages:
        logger.warning("PDF '%s' yielded no extractable text (possibly image-only).", source)
        return ""

    full_text = "\n\n".join(pages)
    logger.info("PDF '%s': %d chars from %d pages", source, len(full_text), len(pages))
    return full_text


def _read_docx_bytes(data: bytes, source: str) -> str:
    try:
        import docx
        import io
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = docx.Document(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts)
    logger.info("DOCX '%s': %d chars", source, len(full_text))
    return full_text


def _extract_html_text(html: str, url: str) -> str:
    """Extract readable text from an HTML page.

    Strategy:
      1. Try html2text (best quality, preserves structure)
      2. Fall back to BeautifulSoup (always available via httpx deps)
      3. Last resort: regex strip
    """
    # Strategy 1: html2text
    try:
        import html2text
        h = html2text.HTML2Text()
        h.ignore_links = True
        h.ignore_images = True
        h.ignore_emphasis = False
        h.body_width = 0   # no wrapping
        text = h.handle(html)
        text = _clean_text(text)
        logger.info("HTML '%s': extracted %d chars via html2text", url, len(text))
        return text
    except ImportError:
        pass

    # Strategy 2: BeautifulSoup
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        # Remove script/style noise
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        text = _clean_text(text)
        logger.info("HTML '%s': extracted %d chars via BeautifulSoup", url, len(text))
        return text
    except ImportError:
        pass

    # Strategy 3: regex fallback
    text = re.sub(r"<[^>]+>", " ", html)
    text = _clean_text(text)
    logger.warning("HTML '%s': extracted %d chars via regex (install html2text for better quality)", url, len(text))
    return text


def _clean_text(text: str) -> str:
    """Collapse excessive whitespace."""
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


# ── Format detection helper ────────────────────────────────────────


def detect_cv_format(cv_path: Path) -> str:
    """Return a human-readable format label for logging."""
    suffix = Path(cv_path).suffix.lower()
    mapping = {
        ".md": "Markdown",
        ".txt": "Plain text",
        ".pdf": "PDF",
        ".docx": "Word (docx)",
        ".doc": "Word (doc)",
    }
    return mapping.get(suffix, f"Unknown ({suffix})")
