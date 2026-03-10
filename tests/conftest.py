"""pytest configuration and shared fixtures.

Automatically generates binary test fixtures (PDF, DOCX) at test time
so no binary files need to be committed to the repository.
"""

from __future__ import annotations

import io
import json
from pathlib import Path
from unittest.mock import MagicMock

import pytest

FIXTURES = Path(__file__).parent / "fixtures"

# ── CV content shared across all format fixtures ───────────────────

CV_CONTENT = """\
John Doe
AI Researcher & Engineer | Hannover, Germany
Email: john.doe@example.com
LinkedIn: https://linkedin.com/in/johndoe
GitHub: https://github.com/johndoe

SUMMARY
PhD candidate in Computer Science at Leibniz University Hannover, specializing in
Knowledge Graphs, Causal Inference, and Large Language Models. 3+ years of research
experience with publications in top venues.

EDUCATION
PhD in Computer Science - Leibniz University Hannover (2022-present)
Thesis: Causality with Knowledge Graphs: Semantics and Inference

M.Sc. in Computer Science - Technical University Munich (2019-2022), GPA: 1.3

EXPERIENCE
Research Assistant - L3S Research Center, Hannover (2022-present)
- Developed KG-enhanced causal discovery algorithms
- Published 3 papers on neuro-symbolic AI
- Built RAG pipelines with LLMs for knowledge extraction
- Technologies: Python, PyTorch, SPARQL, LangChain

ML Engineer Intern - Siemens AI Lab, Munich (2021-2022)
- Implemented NLP pipeline for industrial document processing
- Deployed models on edge devices (Jetson)
- Technologies: Python, TensorFlow, Docker, Kubernetes

SKILLS
Technical: Python, PyTorch, TensorFlow, SPARQL, SQL, Docker, Linux
Frameworks: LangChain, FastAPI, pgmpy, NetworkX, React
Domains: NLP, Causal Discovery, RAG, Knowledge Graphs, Edge AI
Languages: English (fluent), German (B2), Chinese (native)

PUBLICATIONS
KG-Enhanced Causal Discovery - IEEE Access, 2024
Causality with Knowledge Graphs - WSDM Doctoral Consortium, 2026
"""


# ── Session-scoped fixture generators ─────────────────────────────

@pytest.fixture(scope="session")
def sample_pdf_path() -> Path:
    """Generate a real PDF fixture using reportlab (if available) and return path."""
    out = FIXTURES / "sample_cv.pdf"
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        c.setFont("Helvetica", 11)
        y = 800
        for line in CV_CONTENT.splitlines():
            c.drawString(50, y, line)
            y -= 15
            if y < 50:
                c.showPage()
                y = 800
        c.save()
        out.write_bytes(buf.getvalue())
    except ImportError:
        # Fallback: minimal valid PDF (text-searchable via pypdf)
        _write_minimal_pdf(out, CV_CONTENT)
    return out


@pytest.fixture(scope="session")
def sample_docx_path() -> Path:
    """Generate a real DOCX fixture using python-docx and return path."""
    import docx as python_docx

    out = FIXTURES / "sample_cv.docx"
    doc = python_docx.Document()
    doc.add_heading("John Doe", 0)
    doc.add_paragraph("AI Researcher & Engineer | Hannover, Germany")
    doc.add_paragraph("Email: john.doe@example.com")
    doc.add_paragraph("LinkedIn: https://linkedin.com/in/johndoe")

    doc.add_heading("Summary", 1)
    doc.add_paragraph(
        "PhD candidate in Computer Science at Leibniz University Hannover, "
        "specializing in Knowledge Graphs, Causal Inference, and Large Language Models."
    )

    doc.add_heading("Education", 1)
    doc.add_paragraph("PhD in Computer Science - Leibniz University Hannover (2022-present)")
    doc.add_paragraph("Thesis: Causality with Knowledge Graphs: Semantics and Inference")
    doc.add_paragraph("M.Sc. in Computer Science - Technical University Munich (2019-2022)")

    doc.add_heading("Experience", 1)
    doc.add_paragraph("Research Assistant - L3S Research Center, Hannover (2022-present)")
    for item in [
        "Developed KG-enhanced causal discovery algorithms",
        "Published 3 papers on neuro-symbolic AI",
        "Technologies: Python, PyTorch, SPARQL, LangChain",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("ML Engineer Intern - Siemens AI Lab, Munich (2021-2022)")
    for item in [
        "Implemented NLP pipeline for industrial document processing",
        "Technologies: Python, TensorFlow, Docker, Kubernetes",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Skills", 1)
    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "Technical"
    t.cell(0, 1).text = "Python, PyTorch, TensorFlow, SPARQL, SQL, Docker, Linux"
    t.cell(1, 0).text = "Frameworks"
    t.cell(1, 1).text = "LangChain, FastAPI, pgmpy, NetworkX, React"
    t.cell(2, 0).text = "Domains"
    t.cell(2, 1).text = "NLP, Causal Discovery, RAG, Knowledge Graphs, Edge AI"

    doc.add_heading("Publications", 1)
    doc.add_paragraph("KG-Enhanced Causal Discovery - IEEE Access, 2024")
    doc.add_paragraph("Causality with Knowledge Graphs - WSDM Doctoral Consortium, 2026")

    doc.save(str(out))
    return out


@pytest.fixture
def mock_llm(sample_profile_data):
    """A mock LLM client that returns the sample profile JSON."""
    m = MagicMock()
    m.endpoint.provider = "openai"
    m.endpoint.model = "test-model"
    m.complete_json.return_value = sample_profile_data
    m.complete_structured.return_value = sample_profile_data
    return m


@pytest.fixture(scope="session")
def sample_profile_data():
    return json.loads((FIXTURES / "sample_profile.json").read_text())


# ── Minimal PDF writer (no reportlab) ────────────────────────────

def _write_minimal_pdf(path: Path, text: str) -> None:
    """Write a minimal but pypdf-parseable PDF with embedded text."""
    lines = text.splitlines()
    content_parts = ["BT", "/F1 10 Tf", "50 800 Td", "12 TL"]
    for line in lines:
        safe = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        content_parts.append(f"({safe}) Tj T*")
    content_parts.append("ET")
    content = "\n".join(content_parts)
    stream = content.encode("latin-1")

    pdf = (
        f"%PDF-1.4\n"
        f"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        f"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        f"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        f"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        f"4 0 obj<</Length {len(stream)}>>\nstream\n"
        f"{content}\nendstream\nendobj\n"
        f"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        f"xref\n0 6\n"
        f"0000000000 65535 f \n"
        f"trailer<</Size 6/Root 1 0 R>>\nstartxref\n9\n%%EOF\n"
    )
    path.write_bytes(pdf.encode("latin-1"))
